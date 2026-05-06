from __future__ import annotations

import base64
import io
import json
import math
import os
import re
import statistics
import uuid
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pypdf import PdfReader
from scipy import stats

try:
    from openai import OpenAI
except Exception:  # pragma: no cover
    OpenAI = None

APP_ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = APP_ROOT / 'outputs'
OUTPUT_DIR.mkdir(exist_ok=True)

app = FastAPI(title='SpineResearch Studio Pro', version='1.0.0')
app.add_middleware(
    CORSMiddleware,
    allow_origins=['*'],
    allow_methods=['*'],
    allow_headers=['*'],
)

WORKSPACES: Dict[str, Dict[str, Any]] = {}
DATASETS: Dict[str, pd.DataFrame] = {}

SPINE_TERMS = [
    'spine', 'cervical', 'lumbar', 'thoracic', 'deformity', 'scoliosis', 'ais',
    'fusion', 'tlif', 'acdf', 'laminectomy', 'laminoplasty', 'promis', 'odi', 'ndi', 'vas', 'nprs',
    'mjoa', 'lordosis', 'pelvic', 'sagittal', 'infection', 'osteomyelitis', 'pseudoarthrosis'
]

JOURNAL_OPTIONS = [
    'Clinical Spine Surgery', 'Spine', 'The Spine Journal', 'Global Spine Journal',
    'European Spine Journal', 'Journal of Neurosurgery: Spine', 'Spine Deformity',
    'North American Spine Society Journal', 'World Neurosurgery', 'Neurosurgery', 'Other / not sure'
]

STUDY_DESIGNS = [
    'Not sure / infer from files', 'Retrospective cohort', 'Prospective cohort', 'Case-control',
    'Cross-sectional', 'Systematic review/meta-analysis', 'Prediction model', 'Quality improvement',
    'Randomized trial', 'Other'
]


def safe_float(x: Any) -> Optional[float]:
    try:
        if x is None or (isinstance(x, float) and math.isnan(x)):
            return None
        s = str(x).strip().replace('%', '')
        if s == '' or s.lower() in {'nan', 'none', 'null', 'na', 'n/a'}:
            return None
        return float(s)
    except Exception:
        return None


def normalize_col(name: str) -> str:
    return re.sub(r'[^a-z0-9]+', '_', str(name).lower()).strip('_')


def clean_df(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [str(c).strip() if str(c).strip() else f'unnamed_{i}' for i, c in enumerate(df.columns)]
    df = df.dropna(axis=0, how='all').dropna(axis=1, how='all')
    # Convert object columns that are mostly numeric
    for col in df.columns:
        if df[col].dtype == 'object':
            converted = pd.to_numeric(df[col].astype(str).str.replace('%', '', regex=False).str.strip(), errors='coerce')
            if len(df[col]) and converted.notna().mean() >= 0.75:
                df[col] = converted
    return df


def parse_csv_bytes(data: bytes, name: str) -> pd.DataFrame:
    encodings = ['utf-8-sig', 'utf-8', 'latin-1']
    last_err = None
    for enc in encodings:
        try:
            return clean_df(pd.read_csv(io.BytesIO(data), encoding=enc))
        except Exception as e:
            last_err = e
    raise ValueError(f'Could not parse CSV {name}: {last_err}')


def parse_excel_bytes(data: bytes, name: str) -> pd.DataFrame:
    xls = pd.ExcelFile(io.BytesIO(data))
    # Select largest non-empty sheet
    best_df = None
    best_sheet = None
    best_cells = -1
    for sheet in xls.sheet_names:
        try:
            df = pd.read_excel(io.BytesIO(data), sheet_name=sheet)
            df = clean_df(df)
            cells = df.shape[0] * df.shape[1]
            if cells > best_cells:
                best_df, best_sheet, best_cells = df, sheet, cells
        except Exception:
            continue
    if best_df is None:
        raise ValueError(f'Could not parse Excel file {name}')
    best_df.attrs['sheet_name'] = best_sheet
    return best_df


def parse_docx_bytes(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])


def parse_pdf_bytes(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages[:100]:
        pages.append(page.extract_text() or '')
    return '\n'.join(pages).strip()


def parse_uploaded_file(name: str, data: bytes) -> Tuple[List[Dict[str, Any]], List[Tuple[str, pd.DataFrame]], List[str]]:
    lower = name.lower()
    files_info: List[Dict[str, Any]] = []
    datasets: List[Tuple[str, pd.DataFrame]] = []
    texts: List[str] = []

    def add_info(status: str, parsed: bool, notes: str):
        files_info.append({'name': name, 'type': lower.split('.')[-1].upper() if '.' in lower else 'UNKNOWN', 'status': status, 'parsed': parsed, 'notes': notes})

    try:
        if lower.endswith('.csv'):
            df = parse_csv_bytes(data, name)
            datasets.append((name, df))
            add_info('Dataset parsed', True, f'{len(df)} rows, {len(df.columns)} columns')
        elif lower.endswith(('.xlsx', '.xls')):
            df = parse_excel_bytes(data, name)
            sheet = df.attrs.get('sheet_name', 'sheet')
            datasets.append((name, df))
            add_info('Dataset parsed', True, f'{len(df)} rows, {len(df.columns)} columns from {sheet}')
        elif lower.endswith('.docx'):
            text = parse_docx_bytes(data)
            texts.append(f'--- {name} ---\n{text}')
            add_info('Text extracted', True, f'{len(text.split())} words extracted')
        elif lower.endswith('.pdf'):
            text = parse_pdf_bytes(data)
            texts.append(f'--- {name} ---\n{text}')
            add_info('Text extracted', bool(text), f'{len(text.split())} words extracted; verify PDF extraction')
        elif lower.endswith(('.txt', '.md')):
            text = data.decode('utf-8', errors='ignore')
            texts.append(f'--- {name} ---\n{text}')
            add_info('Text extracted', True, f'{len(text.split())} words extracted')
        elif lower.endswith('.json'):
            text = data.decode('utf-8', errors='ignore')
            texts.append(f'--- {name} ---\n{text[:20000]}')
            add_info('Text extracted', True, 'JSON text cataloged')
        elif lower.endswith('.zip'):
            z = zipfile.ZipFile(io.BytesIO(data))
            parsed_count = 0
            for member in z.namelist():
                if member.endswith('/'):
                    continue
                if member.startswith('__MACOSX/'):
                    continue
                mdata = z.read(member)
                sub_info, sub_datasets, sub_texts = parse_uploaded_file(member, mdata)
                files_info.extend([{**i, 'name': f'{name} / {i["name"]}'} for i in sub_info])
                datasets.extend([(f'{name} / {dname}', df) for dname, df in sub_datasets])
                texts.extend(sub_texts)
                parsed_count += len(sub_datasets) + len(sub_texts)
            files_info.append({'name': name, 'type': 'ZIP', 'status': 'Expanded', 'parsed': parsed_count > 0, 'notes': f'{parsed_count} supported items parsed inside ZIP'})
        else:
            add_info('Cataloged only', False, 'Unsupported type. Convert to CSV, XLSX, DOCX, PDF, TXT, MD, JSON, or ZIP.')
    except Exception as e:
        add_info('Parse error', False, str(e))
    return files_info, datasets, texts


def infer_id_column(df: pd.DataFrame) -> Optional[str]:
    names = ['patient_id', 'id', 'record_id', 'study_id', 'mrn', 'subject_id']
    norm_map = {normalize_col(c): c for c in df.columns}
    for n in names:
        if n in norm_map:
            return norm_map[n]
    for c in df.columns:
        cn = normalize_col(c)
        if 'id' in cn and df[c].nunique(dropna=True) >= max(2, len(df) * 0.5):
            return c
    return None


def classify_variables(df: pd.DataFrame) -> Dict[str, List[str]]:
    classes = {k: [] for k in ['demographic', 'surgical', 'baseline_outcome', 'postoperative_outcome', 'anchor_mcid', 'comorbidity', 'complication', 'numeric_other', 'categorical_other']}
    for col in df.columns:
        n = normalize_col(col)
        is_numeric = pd.api.types.is_numeric_dtype(df[col])
        if any(x in n for x in ['age', 'sex', 'gender', 'bmi', 'race', 'ethnic', 'smok', 'asa']):
            classes['demographic'].append(col)
        elif any(x in n for x in ['procedure', 'approach', 'technique', 'mis', 'open', 'fusion', 'levels', 'level', 'operative', 'surgery', 'ebl', 'blood_loss', 'time', 'tlif', 'acdf']):
            classes['surgical'].append(col)
        elif any(x in n for x in ['pre', 'baseline', 'initial']) and any(x in n for x in ['promis', 'odi', 'ndi', 'vas', 'nprs', 'mjoa', 'glucose', 'a1c', 'pain', 'function', 'disability']):
            classes['baseline_outcome'].append(col)
        elif any(x in n for x in ['post', 'follow', 'final', '3m', '6m', '1y', '2y', 'pod']) and any(x in n for x in ['promis', 'odi', 'ndi', 'vas', 'nprs', 'mjoa', 'glucose', 'a1c', 'pain', 'function', 'disability']):
            classes['postoperative_outcome'].append(col)
        elif any(x in n for x in ['satisfaction', 'global', 'anchor', 'improved', 'better', 'mcid']):
            classes['anchor_mcid'].append(col)
        elif any(x in n for x in ['diabetes', 'htn', 'hypertension', 'renal', 'ckd', 'copd', 'charlson', 'cci', 'depression']):
            classes['comorbidity'].append(col)
        elif any(x in n for x in ['complication', 'infection', 'readmission', 'reoperation', 'sequela', 'death', 'mortality']):
            classes['complication'].append(col)
        elif is_numeric:
            classes['numeric_other'].append(col)
        else:
            classes['categorical_other'].append(col)
    return classes


def missingness(df: pd.DataFrame) -> List[Dict[str, Any]]:
    out = []
    n = len(df)
    for col in df.columns:
        miss = int(df[col].isna().sum())
        pct = 100 * miss / n if n else 0
        concern = 'High' if pct >= 30 else 'Moderate' if pct >= 10 else 'Low'
        out.append({'variable': col, 'missing_n': miss, 'missing_pct': round(pct, 1), 'concern': concern})
    return sorted(out, key=lambda x: x['missing_pct'], reverse=True)


def numeric_summary(series: pd.Series) -> Dict[str, Any]:
    s = pd.to_numeric(series, errors='coerce').dropna()
    if len(s) == 0:
        return {'n': 0}
    return {
        'n': int(len(s)),
        'mean': round(float(s.mean()), 3),
        'sd': round(float(s.std(ddof=1)), 3) if len(s) > 1 else None,
        'median': round(float(s.median()), 3),
        'q1': round(float(s.quantile(0.25)), 3),
        'q3': round(float(s.quantile(0.75)), 3),
        'min': round(float(s.min()), 3),
        'max': round(float(s.max()), 3),
    }


def dataset_audit(df: pd.DataFrame, filename: str) -> Dict[str, Any]:
    id_col = infer_id_column(df)
    classes = classify_variables(df)
    num_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
    cat_cols = [c for c in df.columns if c not in num_cols]
    plaus = []
    for col in df.columns:
        n = normalize_col(col)
        vals = pd.to_numeric(df[col], errors='coerce')
        if vals.notna().sum() == 0:
            continue
        if n == 'age' or 'age' in n:
            bad = int(((vals < 0) | (vals > 110)).sum())
            if bad: plaus.append(f'{col}: {bad} age values outside 0–110; verify.')
        if 'bmi' in n:
            bad = int(((vals < 10) | (vals > 80)).sum())
            if bad: plaus.append(f'{col}: {bad} BMI values outside 10–80; verify.')
        if any(x in n for x in ['vas', 'nprs']) and 'promis' not in n:
            bad = int(((vals < 0) | (vals > 10)).sum())
            if bad: plaus.append(f'{col}: {bad} VAS/NPRS values outside 0–10; verify.')
        if 'promis' in n:
            bad = int(((vals < 10) | (vals > 90)).sum())
            if bad: plaus.append(f'{col}: {bad} PROMIS-like values outside 10–90; verify.')
    duplicate_ids = None
    if id_col:
        duplicate_ids = int(df[id_col].duplicated().sum())
    return {
        'filename': filename,
        'rows': int(len(df)),
        'columns': int(len(df.columns)),
        'column_names': list(map(str, df.columns)),
        'patient_id_candidate': id_col,
        'duplicate_id_count': duplicate_ids,
        'numeric_variables': list(map(str, num_cols)),
        'categorical_variables': list(map(str, cat_cols)),
        'variable_classification': classes,
        'missingness': missingness(df),
        'numeric_summary': {c: numeric_summary(df[c]) for c in num_cols[:80]},
        'plausibility_warnings': plaus or ['No major simple plausibility warnings detected.'],
    }


def select_dataset(datasets: List[Tuple[str, pd.DataFrame]]) -> Tuple[str, pd.DataFrame]:
    if not datasets:
        raise HTTPException(status_code=400, detail='No structured dataset was parsed. Upload CSV or Excel for analysis.')
    return max(datasets, key=lambda item: item[1].shape[0] * item[1].shape[1])


def find_terms_in_note(note: str, cols: List[str]) -> Dict[str, List[str]]:
    note_n = normalize_col(note)
    found = {'exposures': [], 'outcomes': []}
    exposure_terms = ['approach', 'technique', 'mis', 'minimally_invasive', 'open', 'procedure', 'group', 'diagnosis', 'levels', 'surgical']
    outcome_terms = ['promis', 'odi', 'ndi', 'vas', 'nprs', 'mjoa', 'glucose', 'a1c', 'pain', 'function', 'disability', 'complication', 'infection']
    for col in cols:
        cn = normalize_col(col)
        if any(t in cn for t in exposure_terms) or any(t in note_n and t in cn for t in exposure_terms):
            found['exposures'].append(col)
        if any(t in cn for t in outcome_terms) or any(t in note_n and t in cn for t in outcome_terms):
            found['outcomes'].append(col)
    return found


def suggest_research(note: str, audit: Dict[str, Any]) -> Dict[str, Any]:
    cols = audit['column_names']
    terms = find_terms_in_note(note, cols)
    classes = audit['variable_classification']
    outcomes = terms['outcomes'] or classes['postoperative_outcome'] + classes['baseline_outcome'] + audit['numeric_variables'][:5]
    exposures = terms['exposures'] or classes['surgical'] + classes['demographic'] + classes['categorical_variables'][:5]
    primary_outcome = outcomes[0] if outcomes else 'primary postoperative outcome'
    primary_exposure = exposures[0] if exposures else 'key clinical or surgical predictor'
    spine_context = 'spine surgery' if any(t in normalize_col(note) for t in SPINE_TERMS) else 'the study cohort'
    q1 = f'Among patients in {spine_context}, is {primary_exposure} associated with {primary_outcome}?'
    q2 = f'What baseline, demographic, or surgical factors are associated with clinically meaningful changes in {primary_outcome}?'
    q3 = f'What are the perioperative or postoperative outcome patterns observed in this cohort?'
    hypo = f'{primary_exposure} will be associated with differences in {primary_outcome}, after accounting for relevant baseline and surgical characteristics when sample size permits.'
    not_feasible = []
    if not classes['anchor_mcid']:
        not_feasible.append('Anchor-based MCID derivation is not feasible unless a satisfaction/global rating anchor is provided.')
    if not classes['postoperative_outcome']:
        not_feasible.append('Longitudinal postoperative outcome claims are limited unless postoperative/follow-up outcome fields are confirmed.')
    if audit['rows'] < 50:
        not_feasible.append('Multivariable modeling may be underpowered depending on outcome events and candidate predictors.')
    return {
        'recommended_primary_question': q1,
        'secondary_questions': [q2, q3],
        'hypothesis': hypo,
        'detected_exposure_candidates': exposures[:12],
        'detected_outcome_candidates': outcomes[:12],
        'not_feasible_or_needs_caution': not_feasible or ['No major feasibility barrier detected from the basic audit, but statistical assumptions still require verification.'],
        'study_design_recommendation': 'Retrospective cohort study' if 'retrospective' in normalize_col(note) else 'Observational cohort design likely, unless user confirms a different design.',
        'unresolved_items': [
            'Confirm inclusion and exclusion criteria.', 'Confirm IRB/ethics status.',
            'Confirm exact follow-up timepoint definitions.', 'Confirm final primary outcome and exposure.',
            'Confirm MCID thresholds if MCID analysis is planned.', 'Verify literature citations before journal submission.'
        ]
    }


def table1(df: pd.DataFrame, group_col: Optional[str] = None) -> List[Dict[str, Any]]:
    rows = []
    cols = list(df.columns)
    if group_col and group_col in cols:
        groups = [g for g in df[group_col].dropna().unique().tolist() if str(g).strip() != ''][:8]
    else:
        groups = []
    preferred = []
    for c in cols:
        cn = normalize_col(c)
        if any(t in cn for t in ['age', 'sex', 'gender', 'bmi', 'smok', 'diabetes', 'asa', 'diagnosis', 'procedure', 'approach', 'technique', 'levels', 'level']):
            preferred.append(c)
    selected = preferred[:30] or cols[:20]
    for c in selected:
        if c == group_col:
            continue
        entry = {'variable': c, 'overall': format_var_summary(df[c])}
        for g in groups:
            entry[str(g)] = format_var_summary(df.loc[df[group_col] == g, c])
        rows.append(entry)
    return rows


def format_var_summary(s: pd.Series) -> str:
    numeric = pd.to_numeric(s, errors='coerce')
    if numeric.notna().mean() >= 0.75 and numeric.notna().sum() > 0:
        sm = numeric_summary(s)
        if sm.get('sd') is None:
            return f"{sm.get('mean')}"
        return f"{sm['mean']} ± {sm['sd']}; median {sm['median']} [{sm['q1']}, {sm['q3']}]; n={sm['n']}"
    else:
        vc = s.dropna().astype(str).value_counts().head(3)
        denom = int(s.dropna().shape[0])
        if denom == 0:
            return 'No non-missing data'
        parts = [f'{idx}: {cnt} ({cnt/denom*100:.1f}%)' for idx, cnt in vc.items()]
        return '; '.join(parts)


def paired_screening(df: pd.DataFrame) -> List[Dict[str, Any]]:
    cols = list(df.columns)
    norm_to_col = {normalize_col(c): c for c in cols}
    pre_tokens = ['preop', 'pre', 'baseline', 'initial']
    post_tokens = ['postop', 'post', 'final', 'followup', 'follow_up', '3m', '6m', '1y', '2y', 'pod1', 'pod2']
    results = []
    used = set()
    for c in cols:
        cn = normalize_col(c)
        if c in used or not pd.api.types.is_numeric_dtype(pd.to_numeric(df[c], errors='coerce')):
            continue
        token = next((t for t in pre_tokens if t in cn), None)
        if not token:
            continue
        base = cn.replace(token, '')
        candidates = []
        for pc in cols:
            pcn = normalize_col(pc)
            if pc == c or not any(t in pcn for t in post_tokens):
                continue
            pbase = pcn
            for t in post_tokens:
                pbase = pbase.replace(t, '')
            similarity = len(set(base.split('_')).intersection(set(pbase.split('_'))))
            if base.strip('_') and (base.strip('_') in pbase or pbase.strip('_') in base or similarity >= 1):
                candidates.append(pc)
        if candidates:
            post = candidates[0]
            a = pd.to_numeric(df[c], errors='coerce')
            b = pd.to_numeric(df[post], errors='coerce')
            paired = pd.DataFrame({'pre': a, 'post': b}).dropna()
            if len(paired) >= 3:
                diff = paired['post'] - paired['pre']
                try:
                    p = stats.ttest_rel(paired['post'], paired['pre']).pvalue if len(paired) > 2 else None
                except Exception:
                    p = None
                results.append({
                    'baseline_variable': c,
                    'followup_variable': post,
                    'paired_n': int(len(paired)),
                    'baseline': format_var_summary(paired['pre']),
                    'followup': format_var_summary(paired['post']),
                    'mean_change': round(float(diff.mean()), 3),
                    'p_value_exploratory': format_p(p),
                    'warning': 'Exploratory paired t-test; verify timepoint definitions and distributional assumptions.'
                })
                used.add(c); used.add(post)
    return results[:25]


def format_p(p: Optional[float]) -> str:
    if p is None or math.isnan(p):
        return 'Not calculated'
    if p < 0.001:
        return '<0.001'
    return f'{p:.3f}'


def group_comparison(df: pd.DataFrame, group_col: str, outcome_col: str) -> Dict[str, Any]:
    if group_col not in df.columns or outcome_col not in df.columns:
        raise HTTPException(status_code=400, detail='Selected group or outcome variable not found.')
    y = pd.to_numeric(df[outcome_col], errors='coerce')
    temp = pd.DataFrame({'group': df[group_col].astype(str), 'outcome': y}).dropna()
    temp = temp[temp['group'].str.strip() != '']
    groups = temp['group'].unique().tolist()
    if len(groups) < 2:
        return {'error': 'At least two groups with non-missing outcome data are required.'}
    if len(groups) > 12:
        top = temp['group'].value_counts().head(12).index.tolist()
        temp = temp[temp['group'].isin(top)]
        groups = top
    summary = []
    for g in groups:
        s = temp.loc[temp['group'] == g, 'outcome']
        sm = numeric_summary(s)
        summary.append({'group': str(g), **sm})
    p = None
    test = 'Not calculated'
    warning = 'Exploratory comparison; verify group definitions, missingness, independence, and distribution before publication.'
    if len(groups) == 2:
        a = temp.loc[temp['group'] == groups[0], 'outcome']
        b = temp.loc[temp['group'] == groups[1], 'outcome']
        if len(a) >= 2 and len(b) >= 2:
            p = stats.ttest_ind(a, b, equal_var=False, nan_policy='omit').pvalue
            test = 'Welch t-test'
        mean_diff = float(b.mean() - a.mean()) if len(a) and len(b) else None
    else:
        arrays = [temp.loc[temp['group'] == g, 'outcome'] for g in groups]
        if all(len(arr) >= 2 for arr in arrays):
            p = stats.f_oneway(*arrays).pvalue
            test = 'One-way ANOVA'
        mean_diff = None
    fig = make_bar_figure(summary, outcome_col)
    return {
        'group_variable': group_col,
        'outcome_variable': outcome_col,
        'test': test,
        'p_value_exploratory': format_p(p),
        'mean_difference_note': 'Mean difference is group 2 minus group 1 for two-group comparisons.' if len(groups) == 2 else 'Mean difference not summarized for >2 groups.',
        'summary': summary,
        'figure_svg': fig,
        'warning': warning,
    }


def make_bar_figure(summary: List[Dict[str, Any]], outcome: str) -> str:
    labels = [s['group'] for s in summary]
    means = [s.get('mean') or 0 for s in summary]
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.bar(labels, means)
    ax.set_ylabel(outcome)
    ax.set_title(f'Mean {outcome} by group')
    ax.tick_params(axis='x', rotation=35)
    fig.tight_layout()
    bio = io.BytesIO()
    fig.savefig(bio, format='svg')
    plt.close(fig)
    return bio.getvalue().decode('utf-8')


def missingness_figure(miss: List[Dict[str, Any]]) -> str:
    top = miss[:20]
    labels = [m['variable'] for m in top][::-1]
    vals = [m['missing_pct'] for m in top][::-1]
    fig, ax = plt.subplots(figsize=(8, max(4, len(labels) * 0.25)))
    ax.barh(labels, vals)
    ax.set_xlabel('Missing (%)')
    ax.set_title('Top variables by missingness')
    fig.tight_layout()
    bio = io.BytesIO()
    fig.savefig(bio, format='svg')
    plt.close(fig)
    return bio.getvalue().decode('utf-8')


def build_analysis_package(df: pd.DataFrame, audit: Dict[str, Any], note: str) -> Dict[str, Any]:
    candidates = suggest_research(note, audit)
    group_candidates = []
    for c in audit['categorical_variables']:
        nunique = df[c].nunique(dropna=True)
        if 2 <= nunique <= 12:
            group_candidates.append(c)
    numeric_candidates = audit['numeric_variables']
    auto_group = (candidates['detected_exposure_candidates'] or group_candidates or [None])[0]
    if auto_group not in df.columns:
        auto_group = group_candidates[0] if group_candidates else None
    auto_outcome = (candidates['detected_outcome_candidates'] or numeric_candidates or [None])[0]
    if auto_outcome not in df.columns:
        auto_outcome = numeric_candidates[0] if numeric_candidates else None
    comparison = None
    if auto_group and auto_outcome and auto_group != auto_outcome:
        try:
            comparison = group_comparison(df, auto_group, auto_outcome)
        except Exception as e:
            comparison = {'error': str(e)}
    return {
        'table1': table1(df, auto_group),
        'paired_screening': paired_screening(df),
        'auto_group_variable': auto_group,
        'auto_outcome_variable': auto_outcome,
        'group_candidates': group_candidates,
        'numeric_candidates': numeric_candidates,
        'primary_comparison': comparison,
        'regression_feasibility': regression_feasibility(df, auto_group, auto_outcome),
        'missingness_figure_svg': missingness_figure(audit['missingness']),
    }


def regression_feasibility(df: pd.DataFrame, group: Optional[str], outcome: Optional[str]) -> Dict[str, Any]:
    if not outcome or outcome not in df.columns:
        return {'status': 'Not assessed', 'reason': 'No numeric outcome selected.'}
    n = int(pd.to_numeric(df[outcome], errors='coerce').notna().sum())
    if n < 50:
        return {'status': 'Caution', 'reason': f'Only {n} non-missing observations for the outcome; multivariable regression may be underpowered.'}
    return {'status': 'Potentially feasible', 'reason': f'{n} non-missing observations; final covariate count should be limited and clinically justified.'}


def html_table(rows: List[Dict[str, Any]]) -> str:
    if not rows:
        return '<p>No rows available.</p>'
    cols = list(dict.fromkeys([k for row in rows for k in row.keys() if k != 'figure_svg']))
    h = '<table><thead><tr>' + ''.join(f'<th>{escape(c)}</th>' for c in cols) + '</tr></thead><tbody>'
    for row in rows:
        h += '<tr>' + ''.join(f'<td>{escape(row.get(c, ""))}</td>' for c in cols) + '</tr>'
    h += '</tbody></table>'
    return h


def escape(x: Any) -> str:
    return str(x).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def generate_manuscript(workspace: Dict[str, Any], direction: Optional[str], use_ai: bool = True) -> Dict[str, Any]:
    note = workspace.get('note', '')
    journal = workspace.get('journal') or 'Target journal not specified'
    audit = workspace['audit']
    research = workspace['research']
    analysis = workspace['analysis']
    comparison = analysis.get('primary_comparison') or {}
    direction_text = direction or research['recommended_primary_question']
    title = make_title(direction_text, audit)
    result_lines = []
    result_lines.append(f'The analytic dataset contained {audit["rows"]} rows and {audit["columns"]} variables.')
    if comparison and not comparison.get('error'):
        result_lines.append(f'The primary exploratory comparison evaluated {comparison["outcome_variable"]} across {comparison["group_variable"]}.')
        result_lines.append(f'The exploratory test was {comparison["test"]} with p={comparison["p_value_exploratory"]}.')
    else:
        result_lines.append('A primary comparison was not completed because a valid exposure/outcome pairing was not available or selected.')
    paired = analysis.get('paired_screening') or []
    if paired:
        result_lines.append(f'Paired baseline-to-follow-up screening identified {len(paired)} potential paired outcome comparisons; timepoint definitions require verification.')
    refs_needed = [
        'Citation supporting the clinical importance of the selected outcome measure.',
        'Citation supporting the selected MCID threshold if MCID is used.',
        'Citation supporting the clinical background of the target spine population.',
        'Citation for prior studies evaluating similar surgical exposure or outcome.',
        'Target journal instructions for abstract structure, word count, and reference style.'
    ]
    sections = {
        'Title': title,
        'Short title': title[:70],
        'Keywords': 'spine surgery; patient-reported outcomes; retrospective cohort; statistical analysis; manuscript',
        'Abstract': build_abstract(title, direction_text, audit, result_lines),
        'Introduction': (
            f'This manuscript is centered on the user-provided study concept: {note.strip() or "[No rough note provided]"}. '
            f'The proposed research question is: {direction_text}. The introduction should establish the clinical importance of the population, '
            'the relevance of the selected outcome or exposure, and the specific knowledge gap addressed by this dataset. '
            '[Literature citations required; none are invented by this application.]'
        ),
        'Methods': (
            f'This appears most consistent with {research.get("study_design_recommendation", "an observational study")}. '
            f'The dataset included {audit["rows"]} records and {audit["columns"]} variables. Variables were classified into demographic, surgical, outcome, anchor/MCID, comorbidity, and complication domains based on column names and data structure. '
            'Final inclusion/exclusion criteria, IRB approval, and follow-up definitions require user confirmation.'
        ),
        'Statistical Analysis': (
            'Continuous variables should be summarized as mean ± standard deviation or median [interquartile range], depending on distribution. '
            'Categorical variables should be summarized as counts and percentages. Exploratory group comparisons were computed when a categorical exposure and numeric outcome were available. '
            'Paired analyses were screened when baseline and postoperative/follow-up variable names appeared to match. Any p-values generated by this prototype should be verified in a formal statistical workflow before submission.'
        ),
        'Results': ' '.join(result_lines),
        'Discussion': (
            'The discussion should interpret the verified findings in the context of surgical invasiveness, patient characteristics, outcomes, and clinical relevance. '
            'Do not claim causality unless the final study design and analytic approach justify causal inference. Comparison with prior literature must be added using verified references.'
        ),
        'Limitations': (
            'Potential limitations include observational design, missing follow-up data, possible selection bias, residual confounding, uncertainty in variable definitions, and need for external validation. '
            'Additional limitations should be tailored after confirming data source, inclusion criteria, follow-up timing, and final statistical models.'
        ),
        'Conclusion': (
            f'The uploaded dataset appears to support a manuscript evaluating: {direction_text}. Conclusions should remain cautious and reflect only verified analyses.'
        ),
        'References needed': '\n'.join(f'- {r}' for r in refs_needed),
        'Unresolved checklist': '\n'.join(f'- {x}' for x in research.get('unresolved_items', [])),
    }

    if use_ai and os.environ.get('OPENAI_API_KEY') and OpenAI:
        ai_text = ai_refine_manuscript(sections, workspace, direction_text)
        if ai_text:
            sections['AI-refined manuscript draft'] = ai_text
    return {'direction': direction_text, 'sections': sections, 'references_needed': refs_needed}


def make_title(direction: str, audit: Dict[str, Any]) -> str:
    d = direction.strip().rstrip('?')
    return d[0].upper() + d[1:] + ': A Data-Grounded Spine Research Analysis'


def build_abstract(title: str, direction: str, audit: Dict[str, Any], result_lines: List[str]) -> str:
    return (
        'Background: Spine surgery research often requires careful alignment between clinical questions, available data, and feasible statistical analysis.\n'
        f'Objective: To evaluate {direction.rstrip("?").lower()}.\n'
        f'Methods: Uploaded files were processed to identify a structured dataset with {audit["rows"]} records and {audit["columns"]} variables. Dataset audit, missingness assessment, variable classification, exploratory Table 1 generation, and selected outcome analyses were performed.\n'
        f'Results: {' '.join(result_lines)}\n'
        'Conclusions: The dataset appears to support a cautious data-grounded manuscript, but final statistical verification, reference validation, and user confirmation of clinical definitions are required before journal submission.'
    )


def ai_refine_manuscript(sections: Dict[str, str], workspace: Dict[str, Any], direction: str) -> Optional[str]:
    try:
        client = OpenAI()
        payload = {
            'direction': direction,
            'audit': workspace['audit'],
            'research': workspace['research'],
            'analysis_summary': {
                'auto_group_variable': workspace['analysis'].get('auto_group_variable'),
                'auto_outcome_variable': workspace['analysis'].get('auto_outcome_variable'),
                'primary_comparison': workspace['analysis'].get('primary_comparison'),
                'paired_screening': workspace['analysis'].get('paired_screening')[:10],
            },
            'draft_sections': sections,
        }
        system = 'You are a careful spine surgery research manuscript assistant. Do not invent citations, sample sizes, p-values, or conclusions. Use only the provided verified data and label missing items clearly.'
        user = 'Create a polished journal-style manuscript draft from this verified analysis payload. Use placeholders only for literature citations and unconfirmed details.\n' + json.dumps(payload)[:50000]
        resp = client.chat.completions.create(model='gpt-4.1-mini', messages=[{'role': 'system', 'content': system}, {'role': 'user', 'content': user}], temperature=0.2)
        return resp.choices[0].message.content
    except Exception:
        return None


def make_docx(manuscript: Dict[str, Any], workspace: Dict[str, Any]) -> Path:
    doc = Document()
    doc.add_heading('SpineResearch Studio Manuscript Package', 0)
    doc.add_paragraph(f'Workspace ID: {workspace["id"]}')
    doc.add_paragraph('Important: This draft is generated from uploaded data and computed exploratory outputs. Verify all statistics, references, and journal requirements before submission.')
    for heading, text in manuscript['sections'].items():
        doc.add_heading(heading, level=1)
        for para in str(text).split('\n'):
            doc.add_paragraph(para)
    doc.add_heading('Table 1: Exploratory baseline/descriptive table', level=1)
    add_table_to_doc(doc, workspace['analysis'].get('table1', [])[:40])
    comp = workspace['analysis'].get('primary_comparison')
    if comp and not comp.get('error'):
        doc.add_heading('Primary comparison summary', level=1)
        add_table_to_doc(doc, comp.get('summary', []))
    paired = workspace['analysis'].get('paired_screening')
    if paired:
        doc.add_heading('Paired outcome screening', level=1)
        add_table_to_doc(doc, paired)
    path = OUTPUT_DIR / f'manuscript_{workspace["id"]}.docx'
    doc.save(path)
    return path


def add_table_to_doc(doc: Document, rows: List[Dict[str, Any]]):
    if not rows:
        doc.add_paragraph('No table rows available.')
        return
    cols = list(dict.fromkeys([k for row in rows for k in row.keys() if k != 'figure_svg']))
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    for i, c in enumerate(cols):
        table.rows[0].cells[i].text = str(c)
    for row in rows:
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            cells[i].text = str(row.get(c, ''))


@app.get('/health')
def health():
    return {'ok': True, 'service': 'SpineResearch Studio Pro'}


@app.post('/api/analyze')
async def analyze(
    note: str = Form(''),
    journal: str = Form(''),
    study_design: str = Form(''),
    files: List[UploadFile] = File(default=[]),
):
    all_info: List[Dict[str, Any]] = []
    datasets: List[Tuple[str, pd.DataFrame]] = []
    texts: List[str] = []
    for f in files:
        data = await f.read()
        info, ds, tx = parse_uploaded_file(f.filename or 'uploaded_file', data)
        all_info.extend(info); datasets.extend(ds); texts.extend(tx)
    if not datasets:
        raise HTTPException(status_code=400, detail='No analyzable structured dataset found. Upload at least one CSV or Excel file for statistical analysis.')
    fname, df = select_dataset(datasets)
    audit = dataset_audit(df, fname)
    research = suggest_research(note + '\n' + '\n'.join(texts)[:15000], audit)
    if study_design and study_design != 'Not sure / infer from files':
        research['study_design_recommendation'] = study_design
    analysis = build_analysis_package(df, audit, note)
    wid = str(uuid.uuid4())[:12]
    DATASETS[wid] = df
    workspace = {
        'id': wid,
        'note': note,
        'journal': journal or 'Not specified',
        'study_design': study_design or 'Not specified',
        'files': all_info,
        'supporting_text_excerpt': '\n\n'.join(texts)[:30000],
        'audit': audit,
        'research': research,
        'analysis': analysis,
    }
    WORKSPACES[wid] = workspace
    # Do not return huge df
    return workspace


@app.post('/api/compare/{workspace_id}')
def compare(workspace_id: str, payload: Dict[str, Any]):
    if workspace_id not in WORKSPACES or workspace_id not in DATASETS:
        raise HTTPException(status_code=404, detail='Workspace not found')
    group_col = payload.get('group_variable')
    outcome_col = payload.get('outcome_variable')
    df = DATASETS[workspace_id]
    comp = group_comparison(df, group_col, outcome_col)
    WORKSPACES[workspace_id]['analysis']['primary_comparison'] = comp
    WORKSPACES[workspace_id]['analysis']['auto_group_variable'] = group_col
    WORKSPACES[workspace_id]['analysis']['auto_outcome_variable'] = outcome_col
    return comp


@app.post('/api/manuscript/{workspace_id}')
def manuscript(workspace_id: str, payload: Dict[str, Any]):
    if workspace_id not in WORKSPACES:
        raise HTTPException(status_code=404, detail='Workspace not found')
    direction = payload.get('direction')
    use_ai = bool(payload.get('use_ai', True))
    ms = generate_manuscript(WORKSPACES[workspace_id], direction, use_ai=use_ai)
    WORKSPACES[workspace_id]['manuscript'] = ms
    return ms


@app.get('/api/download/manuscript/{workspace_id}')
def download_manuscript(workspace_id: str):
    if workspace_id not in WORKSPACES:
        raise HTTPException(status_code=404, detail='Workspace not found')
    ws = WORKSPACES[workspace_id]
    if 'manuscript' not in ws:
        ws['manuscript'] = generate_manuscript(ws, None, use_ai=False)
    path = make_docx(ws['manuscript'], ws)
    return FileResponse(path, filename=f'SpineResearch_Manuscript_{workspace_id}.docx')


@app.get('/api/download/workspace/{workspace_id}')
def download_workspace(workspace_id: str):
    if workspace_id not in WORKSPACES:
        raise HTTPException(status_code=404, detail='Workspace not found')
    path = OUTPUT_DIR / f'workspace_{workspace_id}.json'
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(WORKSPACES[workspace_id], f, ensure_ascii=False, indent=2)
    return FileResponse(path, filename=f'SpineResearch_Workspace_{workspace_id}.json')


app.mount('/', StaticFiles(directory=APP_ROOT / 'static', html=True), name='static')
